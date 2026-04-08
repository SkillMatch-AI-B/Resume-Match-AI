import sys
import os
import hashlib 
import re
import io
import docx 
import textwrap # <--- THE PERMANENT FIX FOR PDF CRASHES
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# --- PDF IMPORT ---
try:
    from fpdf import FPDF
    HAS_FPDF = True
except ImportError:
    HAS_FPDF = False

root_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
if root_dir not in sys.path:
    sys.path.append(root_dir)

import streamlit as st
import database.db_queries as db
import core.resume_parser as parser
import core.nlp_engine as matcher
import core.llm_helper as llm
import utils.report_generator as rg
import core.llm_helper as llm_helper

# ==========================================
# 1. DYNAMIC MARKDOWN GENERATOR
# ==========================================
def generate_markdown_from_json(data):
    if not data: return ""
    md = f"# {data.get('name', 'Name')}\n{data.get('contact', '')}\n\n"
    if data.get('summary', '').strip(): md += f"## PROFESSIONAL SUMMARY\n{data['summary']}\n\n"
    if data.get('skills', '').strip(): md += f"## SKILLS\n{data['skills']}\n\n"
    if data.get('experience', '').strip(): md += f"## EXPERIENCE\n{data['experience']}\n\n"
    if data.get('projects', '').strip(): md += f"## PROJECTS\n{data['projects']}\n\n"
    if data.get('education', '').strip(): md += f"## EDUCATION\n{data['education']}\n\n"
    if data.get('certifications', '').strip(): md += f"## CERTIFICATIONS & ACHIEVEMENTS\n{data['certifications']}\n\n"
    return md

# ==========================================
# 2. DYNAMIC HTML GENERATOR 
# ==========================================
def generate_html_from_json(data, template="Executive ATS"):
    if not data: return ""
    if template == "Modern Tech":
        font_family = "'Segoe UI', Tahoma, Geneva, Verdana, sans-serif"
        h1_align = "left"
        h2_border = "2px solid #2563eb" 
        h2_color = "#1e3a8a"
    elif template == "Minimalist Clean":
        font_family = "'Helvetica Neue', Helvetica, Arial, sans-serif"
        h1_align = "center"
        h2_border = "1px solid #cbd5e1" 
        h2_color = "#334155"
    else: 
        font_family = "'Times New Roman', Times, serif"
        h1_align = "center"
        h2_border = "2px solid black"
        h2_color = "black"

    html = f"<div style='font-family: {font_family}; color: #000;'>"
    html += f"<h1 style='text-align: {h1_align}; color: {h2_color}; font-size: 26px; font-weight: bold; margin-bottom: 4px; text-transform: uppercase;'>{data.get('name', 'Name')}</h1>"
    html += f"<p style='text-align: {h1_align}; margin-bottom: 20px; font-style: italic; color: #475569;'>{data.get('contact', '')}</p>"

    sections = [
        ("PROFESSIONAL SUMMARY", 'summary'), ("SKILLS", 'skills'),
        ("EXPERIENCE", 'experience'), ("PROJECTS", 'projects'),
        ("EDUCATION", 'education'), ("CERTIFICATIONS & ACHIEVEMENTS", 'certifications')
    ]

    for title, key in sections:
        val = data.get(key, '').strip()
        if val:
            val = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', val)
            val = val.replace('\n- ', '<br>• ').replace('\n* ', '<br>• ')
            if val.startswith('- ') or val.startswith('* '): val = '• ' + val[2:]
            val = val.replace('\n', '<br>')
            html += f"<h2 style='border-bottom: {h2_border}; color: {h2_color}; font-size: 16px; margin-top: 15px; margin-bottom: 8px; text-transform: uppercase;'>{title}</h2>"
            html += f"<p style='margin-bottom: 5px; line-height: 1.5;'>{val}</p>"
    html += "</div>"
    return html

# ==========================================
# 3. ADVANCED DOCX FORMATTER 
# ==========================================
def add_horizontal_line(paragraph, r, g, b):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr, 'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap', 'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN', 'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind', 'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc', 'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap', 'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr', 'w:pPrChange')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    hex_color = f"{r:02x}{g:02x}{b:02x}"
    bottom.set(qn('w:color'), hex_color if hex_color != "000000" else "auto")
    pBdr.append(bottom)

def create_formatted_docx(markdown_text, template="Executive ATS"):
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    style = doc.styles['Normal']
    
    if template == "Modern Tech":
        style.font.name = 'Arial'
        hr_r, hr_g, hr_b = 30, 58, 138 
        h1_align = WD_ALIGN_PARAGRAPH.LEFT
    elif template == "Minimalist Clean":
        style.font.name = 'Calibri'
        hr_r, hr_g, hr_b = 51, 65, 85 
        h1_align = WD_ALIGN_PARAGRAPH.CENTER
    else: 
        style.font.name = 'Times New Roman'
        hr_r, hr_g, hr_b = 0, 0, 0 
        h1_align = WD_ALIGN_PARAGRAPH.CENTER

    style.font.size = Pt(11)

    for line in markdown_text.split('\n'):
        line = line.strip()
        if not line: continue
            
        if line.startswith('#'):
            level = line.count('#')
            clean_text = line.replace('#', '').strip().replace('**', '') 
            
            if level == 1:
                p = doc.add_paragraph()
                p.alignment = h1_align
                run = p.add_run(clean_text)
                run.bold = True
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(hr_r, hr_g, hr_b)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(clean_text.upper())
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(hr_r, hr_g, hr_b) 
                add_horizontal_line(p, hr_r, hr_g, hr_b) 
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2) 
            p.paragraph_format.left_indent = Inches(0.25)
            parts = re.split(r'\*\*(.*?)\*\*', line[2:])
            for i, part in enumerate(parts):
                if i % 2 == 1: p.add_run(part).bold = True 
                else: p.add_run(part) 
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            if " | " in line and len(line) < 100:
                p.alignment = h1_align
            parts = re.split(r'\*\*(.*?)\*\*', line)
            for i, part in enumerate(parts):
                if i % 2 == 1: p.add_run(part).bold = True
                else: p.add_run(part)
                
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# ==========================================
# 4. RESUME PDF FORMATTER (BULLETPROOF TEXTWRAP)
# ==========================================
def create_resume_pdf(data, template="Executive ATS"):
    if not HAS_FPDF: return b"Error: FPDF not installed."
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    if template == "Modern Tech":
        title_font = "Arial"
        hr_r, hr_g, hr_b = 30, 58, 138 
        align = 'L'
    elif template == "Minimalist Clean":
        title_font = "Arial"
        hr_r, hr_g, hr_b = 51, 65, 85 
        align = 'C'
    else: 
        title_font = "Times"
        hr_r, hr_g, hr_b = 0, 0, 0 
        align = 'C'

    safe_name = str(data.get('name') or 'Name').strip()
    safe_contact = str(data.get('contact') or '').replace('|', ' | ').strip()

    pdf.set_font(title_font, 'B', 16)
    pdf.set_text_color(hr_r, hr_g, hr_b)
    # Using cell and truncation to avoid multi_cell crashes on headers
    pdf.cell(0, 8, safe_name[:80], ln=True, align=align)
    
    if safe_contact:
        pdf.set_font(title_font, 'I', 10)
        pdf.set_text_color(71, 85, 105) 
        # Textwrap scissors to prevent unbroken URL crashes
        wrapped_contact = textwrap.wrap(safe_contact, width=95, break_long_words=True)
        for wc in wrapped_contact:
            pdf.cell(0, 6, wc, ln=True, align=align)
    pdf.ln(5)

    sections = [
        ("PROFESSIONAL SUMMARY", 'summary'), ("SKILLS", 'skills'),
        ("EXPERIENCE", 'experience'), ("PROJECTS", 'projects'),
        ("EDUCATION", 'education'), ("CERTIFICATIONS & ACHIEVEMENTS", 'certifications')
    ]

    for title, key in sections:
        val = str(data.get(key) or '').strip()
        if val:
            pdf.set_font(title_font, 'B', 12)
            pdf.set_text_color(hr_r, hr_g, hr_b)
            pdf.cell(0, 8, title, ln=True)
            
            pdf.set_draw_color(hr_r, hr_g, hr_b)
            pdf.set_line_width(0.3)
            pdf.line(pdf.get_x(), pdf.get_y(), 210 - pdf.get_x(), pdf.get_y())
            pdf.ln(2)
            
            pdf.set_font(title_font, '', 10)
            pdf.set_text_color(0, 0, 0)
            clean_val = val.replace('**', '').replace('* ', '- ') 
            
            for line in clean_val.split('\n'):
                safe_text = line.encode('latin-1', 'replace').decode('latin-1')
                if not safe_text.strip():
                    pdf.ln(4)
                    continue
                
                # THE FIX: textwrap manually slices any string that is too long
                wrapped_lines = textwrap.wrap(safe_text, width=95, break_long_words=True)
                for w_line in wrapped_lines:
                    pdf.cell(0, 5, w_line, ln=True)
            pdf.ln(4)
            
    try:
        return bytes(pdf.output())
    except Exception:
        return pdf.output(dest='S').encode('latin-1', 'replace')

# ==========================================
# SECURITY CHECK & CONFIG
# ==========================================
st.set_page_config(page_title="AI Resume Builder", page_icon="⚡", layout="wide")

if "logged_in" not in st.session_state or not st.session_state.logged_in:
    st.switch_page("app.py")

user_id = st.session_state.user_data['id']

# ==========================================
# CUSTOM CSS 
# ==========================================
st.markdown("""
    <style>
    .stTextArea textarea, .stTextInput input { font-family: 'Courier New', monospace; font-size: 13px; }
    .streamlit-expanderHeader { font-weight: bold; font-size: 16px; color: #1e293b; }
    .resume-paper {
        background-color: white; color: black;
        padding: 40px; border-radius: 4px;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.1);
        min-height: 800px; font-size: 14px; border: 1px solid #ddd;
    }
    </style>
""", unsafe_allow_html=True)

st.title("⚡ AI Resume Builder")

tab_analyze, tab_history = st.tabs(["🔍 Analysis & Editor", "📚 My Reports"])

# ------------------------------------------
# TAB 1: NEW ANALYSIS
# ------------------------------------------
with tab_analyze:
    if st.session_state.get("show_rewrite_proposal", False):
        st.success(f"Analysis Complete! You scored an **{st.session_state.temp_match_score}%** semantic match.")
        st.markdown("---")
        st.write("### ✨ AI Resume Formatting")
        st.write("Because your resume is a strong match (>75%), would you like our AI to automatically restructure and expertly rewrite it into a professional format?")
        
        c1, c2 = st.columns(2)
        with c1:
            if st.button("Yes, Format & Rewrite Resume ✨", type="primary", use_container_width=True):
                with st.spinner("AI is rewriting grammar and phrasing..."):
                    new_json = llm_helper.rewrite_resume(
                        st.session_state.temp_resume_text, 
                        st.session_state.temp_missing_skills, 
                        st.session_state.temp_job_title
                    )
                    st.session_state.rewritten_resume_json = new_json
                    
                    st.session_state.edit_name = new_json.get('name', '')
                    st.session_state.edit_contact = new_json.get('contact', '')
                    st.session_state.edit_summary = new_json.get('summary', '')
                    st.session_state.edit_skills = new_json.get('skills', '')
                    st.session_state.edit_experience = new_json.get('experience', '')
                    st.session_state.edit_projects = new_json.get('projects', '')
                    st.session_state.edit_education = new_json.get('education', '')
                    st.session_state.edit_certifications = new_json.get('certifications', '')
                    
                    st.session_state.show_rewrite_proposal = False
                    st.rerun()
        with c2:
            if st.button("No, skip for now", use_container_width=True):
                st.session_state.show_rewrite_proposal = False
                st.rerun()

    # ==========================================
    # EDITOR UI & DOWNLOADS
    # ==========================================
    elif st.session_state.get("rewritten_resume_json", None):
        st.markdown("---")
        
        selected_template = st.selectbox("🎨 Choose Resume Template Design", ["Executive ATS", "Modern Tech", "Minimalist Clean"])
        st.markdown("---")

        editor_col, preview_col = st.columns([1, 1], gap="large")
        
        with editor_col:
            st.markdown("### 📝 Live Editor")
            st.info("💡 **Tip:** Click outside the text box after typing to update the preview and downloads!")
            
            with st.expander("👤 Personal Information", expanded=True):
                st.text_input("FULL NAME", key='edit_name')
                st.text_input("CONTACT DETAILS", key='edit_contact')
                
            with st.expander("📝 Professional Summary"):
                st.text_area("SUMMARY", key='edit_summary', height=150)
                
            with st.expander("🛠️ Skills & Categories"):
                st.text_area("SKILLS", key='edit_skills', height=150)
                
            with st.expander("💼 Experience"):
                st.text_area("WORK EXPERIENCE", key='edit_experience', height=200)
                
            with st.expander("🚀 Projects"):
                st.text_area("PROJECTS", key='edit_projects', height=200)
                
            with st.expander("🎓 Education & Certifications"):
                st.text_area("EDUCATION", key='edit_education', height=100)
                st.text_area("CERTIFICATIONS", key='edit_certifications', height=100)

        live_data = {
            'name': st.session_state.edit_name,
            'contact': st.session_state.edit_contact,
            'summary': st.session_state.edit_summary,
            'skills': st.session_state.edit_skills,
            'experience': st.session_state.edit_experience,
            'projects': st.session_state.edit_projects,
            'education': st.session_state.edit_education,
            'certifications': st.session_state.edit_certifications,
        }
            
        with preview_col:
            st.markdown("### 👁️ Document Preview")
            html_preview = generate_html_from_json(live_data, template=selected_template)
            with st.container(height=800, border=False):
                st.markdown(f'<div class="resume-paper">{html_preview}</div>', unsafe_allow_html=True)

        st.markdown("---")
        
        final_markdown = generate_markdown_from_json(live_data)
        
        # --- THE SAFETY SHIELD FOR BUTTONS ---
        btn1, btn2, btn3 = st.columns(3)
        with btn1:
            docx_bytes = create_formatted_docx(final_markdown, template=selected_template)
            st.download_button(label="📄 Download Word (.docx)", data=docx_bytes, file_name="AI_Improved_Resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        with btn2:
            if HAS_FPDF:
                try:
                    pdf_bytes = create_resume_pdf(live_data, template=selected_template)
                    st.download_button(label="📄 Download PDF (.pdf)", data=pdf_bytes, file_name="AI_Improved_Resume.pdf", mime="application/pdf", use_container_width=True)
                except Exception as e:
                    st.error("⚠️ PDF Engine Overloaded. Please use Word (.docx)")
            else:
                st.error("Install 'fpdf2' to enable PDF downloads.")
        with btn3:
            # THIS BUTTON WILL NOW ALWAYS RENDER SAFELY
            if st.button("Close Editor & Start New", type="primary", use_container_width=True):
                st.session_state.rewritten_resume_json = None
                st.rerun()

    else:
        col1, col2 = st.columns(2)
        with col1:
            job_title_input = st.text_input("Target Job Title")
            st.markdown("### Step 1: Provide Your Resume")
            upload_method = st.radio("Choose input method:", ["📁 Upload File", "📝 Paste Text"], horizontal=True)
            resume_text = ""
            resume_name = "Pasted_Text"
            if upload_method == "📁 Upload File":
                resume_file = st.file_uploader("Upload Resume (PDF, DOCX)", type=["pdf", "docx", "txt"])
                if resume_file:
                    resume_name = resume_file.name
                    parsed_result = parser.parse_resume(resume_file)
                    if not parsed_result.startswith("Error:"): resume_text = parsed_result
            else:
                resume_text = st.text_area("Paste text here:", height=150)
        with col2:
            jd_text = st.text_area("Target Job Description:", height=200)

        if st.button("Analyze Resume", type="primary", use_container_width=True):
            if resume_text and jd_text and job_title_input:
                content_hash = hashlib.md5(resume_text.encode('utf-8')).hexdigest()
                
                if db.check_existing_report(user_id, job_title_input, content_hash):
                    st.warning(f"⚠️ You have already generated an analysis for '{job_title_input}' using this exact resume content. Please check your 'My Reports' tab!")
                else:
                    with st.spinner("Processing..."):
                        match_score, _, found_skills = matcher.analyze(resume_text, jd_text)
                        missing_skills = llm_helper.analyze_missing_skills(resume_text, jd_text)
                        ai_feedback = llm.generate_custom_feedback(job_title_input, match_score, missing_skills, resume_text, jd_text)
                        db.save_report(user_id, resume_name, job_title_input, match_score, missing_skills, ai_feedback, content_hash)
                        
                        if match_score > 75:
                            st.session_state.show_rewrite_proposal = True
                            st.session_state.temp_resume_text = resume_text
                            st.session_state.temp_missing_skills = missing_skills
                            st.session_state.temp_job_title = job_title_input
                            st.session_state.temp_match_score = match_score
                            st.rerun() 
                        else:
                            st.success(f"Analysis Complete! You scored an **{match_score}%** semantic match.")
                            st.info("Your full report has been saved to the 'My Reports' tab.")
            else:
                st.warning("Please fill all fields.")

# ------------------------------------------
# TAB 2: MY REPORTS (HISTORY)
# ------------------------------------------
with tab_history:
    reports = db.get_user_reports(user_id)
    if not reports: st.info("No reports yet.")
    else:
        for report in reports:
            with st.expander(f"{'🟢' if report['match_score'] >= 80 else '🟡' if report['match_score'] >= 50 else '🔴'} {report['job_title']} | {report['match_score']}% Match"):
                st.markdown("#### 🎯 Critical Missing Skills")
                st.error(report['missing_skills']) 
                st.markdown("#### 💡 AI Coach Insights")
                st.info(report['ai_feedback']) 
                pdf_bytes = rg.generate_pdf_bytes(
                    job_title=report['job_title'], 
                    match_score=report['match_score'], 
                    missing_skills=report['missing_skills'], 
                    ai_feedback=report['ai_feedback']
                )
                st.markdown("---")
                btn_col1, btn_col2 = st.columns(2)
                with btn_col1:
                    st.download_button(label="📄 Download Full Analysis PDF", data=pdf_bytes, file_name=f"{report['job_title'].replace(' ', '_')}_Report.pdf", mime="application/pdf", key=f"dl_{report['id']}", use_container_width=True)
                with btn_col2:
                    if st.button("🗑️ Delete Report", key=f"delete_{report['id']}", use_container_width=True):
                        db.delete_report(report['id'])
                        st.rerun()