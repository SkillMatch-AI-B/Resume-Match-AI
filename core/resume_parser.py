import pdfplumber
import docx
import os
import re

# def clean_text(raw_text):
#     """
#     Strips out weird hidden characters, corrupted symbols, 
#     and normalizes spacing so the AI can actually read it.
#     """
#     if not raw_text:
#         return ""
#     # Remove non-ASCII characters (often caused by bad fonts or icons)
#     text = re.sub(r'[^\x00-\x7F]+', ' ', raw_text)
#     # Replace multiple spaces, tabs, or newlines with a single space
#     text = re.sub(r'\s+', ' ', text)
#     return text.strip()

def parse_resume(file):
    """Extracts text from PDF, DOCX, or TXT files, including tables."""
    text = ""
    try:
        file_extension = os.path.splitext(file.name)[1].lower()

        if file_extension == '.pdf':
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

        elif file_extension == '.docx':
            doc = docx.Document(file)
            
            # 1. Read standard paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # 2. Read text inside tables (CRUCIAL for this resume)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

        elif file_extension == '.txt':
            text = file.read().decode("utf-8")

        if not text.strip():
            return "Error: The file is empty or could not be read."

        return text.strip()

    except Exception as e:
        return f"Error: {str(e)}"